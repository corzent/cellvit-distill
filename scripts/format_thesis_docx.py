#!/usr/bin/env python3
"""Post-process pandoc-generated thesis.docx to match formal Russian thesis
formatting requirements:

  - A4, margins: left 30mm, right 10mm, top/bottom 20mm
  - Times New Roman, 14pt main text
  - Line spacing 1.5
  - Paragraph first-line indent 1.25 cm
  - Page numbers in footer, centered, arabic; title page counted but not shown
  - Headings (Глава N, sub-sections) kept bold, but same font family
  - Figure and table captions: 12pt italic, centered
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm, Emu
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "Times New Roman"
MAIN_SIZE = Pt(14)
CAPTION_SIZE = Pt(12)
LINE_SPACING = 1.5
FIRST_LINE_INDENT = Cm(1.25)


def set_cell_font(run, font=FONT_NAME, size=MAIN_SIZE, bold=None, italic=None):
    run.font.name = font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_page_number_field(paragraph):
    """Insert PAGE field for page numbering into a paragraph."""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    r = run._element
    r.append(fldChar_begin)
    r.append(instrText)
    r.append(fldChar_sep)
    r.append(fldChar_end)
    set_cell_font(run, size=MAIN_SIZE)


def configure_sections(doc):
    """Set A4 page, margins, and different-first-page header/footer."""
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        # Different first page (title) with no page number
        section.different_first_page_header_footer = True

        # Centered page number in main footer
        footer = section.footer
        # Clear existing content
        for p in list(footer.paragraphs):
            # keep one paragraph, clear content
            for run in list(p.runs):
                run._element.getparent().remove(run._element)
        if not footer.paragraphs:
            footer.add_paragraph()
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_page_number_field(p)

        # Ensure first-page footer is empty (no page number on title)
        first_footer = section.first_page_footer
        for pp in list(first_footer.paragraphs):
            for run in list(pp.runs):
                run._element.getparent().remove(run._element)


def style_paragraph(paragraph, is_heading=False, is_caption=False, indent_first=True):
    """Apply Times New Roman, proper size, line spacing, indent, alignment."""
    # Line spacing
    pf = paragraph.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(6 if is_heading else 0)

    # First-line indent (only for body text, not headings/captions/centered)
    if is_heading or is_caption:
        pf.first_line_indent = Cm(0)
    elif indent_first and paragraph.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
        pf.first_line_indent = FIRST_LINE_INDENT

    # Fonts on all runs
    size = CAPTION_SIZE if is_caption else MAIN_SIZE
    for run in paragraph.runs:
        set_cell_font(run, size=size, italic=True if is_caption else None)

    # Alignment defaults
    if not is_heading and not is_caption and paragraph.alignment is None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def is_caption_paragraph(paragraph):
    """Heuristic: paragraph starting with italic 'Рисунок' or 'Таблица'."""
    text = paragraph.text.strip()
    return text.startswith(("Рисунок ", "Таблица "))


def format_document(path_in, path_out):
    doc = Document(str(path_in))

    # Section / page / margin / footer
    configure_sections(doc)

    # Paragraphs
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading") or style_name.startswith("Title")
        is_caption = is_caption_paragraph(para)

        style_paragraph(para, is_heading=is_heading, is_caption=is_caption)

        # Headings: 14pt bold for H1, slightly smaller for sub-headings
        if is_heading:
            level = 1
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except Exception:
                    level = 1
            heading_size = {1: Pt(16), 2: Pt(15), 3: Pt(14)}.get(level, Pt(14))
            for run in para.runs:
                set_cell_font(run, size=heading_size, bold=True)
            # Headings left-aligned (ГОСТ)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    style_paragraph(para, indent_first=False)
                    for run in para.runs:
                        set_cell_font(run, size=Pt(12))  # tables 12pt for compactness

    doc.save(str(path_out))
    print(f"Saved: {path_out}")


def main():
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("/home/corzent/caspian/thesis/presentation/thesis.docx")
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src.parent / "thesis_formatted.docx"
    format_document(src, dst)


if __name__ == "__main__":
    main()
