#!/usr/bin/env python3
"""Surgical edits to /home/corzent/caspian/thesis/presentation/thesis.docx.

Preserves the user's manually-edited title page (paragraphs 0-17), styles,
margins, and existing tables/images. Applies the following targeted updates
to bring the docx in sync with the latest thesis.md:

  1.  §1.2: insert tissue×class heterogeneity paragraph + heatmap figure 1.4
  2.  §1.4: append NuLite, HoVer-UNet, Mamba/SSM paragraphs
  3.  §2.1: insert KD-variants intro paragraph + pipeline figure 2.1
  4.  §2.3: rewrite student arch (ConvNeXt-Tiny → FastViT-S12)
  5.  §2.4: replace with new intro + KD-variants figure 2.2 + 3 subsections
  6.  §2.6: insert NEW Mamba decoder section (with figure 2.4)
  7.  §2.7: rewrite design-of-experiments (was §2.6) → ablation grid 2×3×3
  8.  Introduction → Научная новизна: rewrite under 4 contributions
  9.  Introduction → Практическая значимость: refined with 3-fold CV numbers

The script is idempotent on its 9 ops: it looks for landmark text or the
heading text exactly and skips an op if its update marker is already present.
"""

import sys
import shutil
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DOCX_PATH = Path("/home/corzent/caspian/thesis/presentation/thesis.docx")
BACKUP_PATH = DOCX_PATH.with_suffix(".bak.docx")
FIG_DIR = Path("/home/corzent/caspian/thesis/thesis_figures")

# Magic marker so re-runs of the script are idempotent.
MARKER = "[v2-sync]"


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def insert_paragraph_after(paragraph, text="", style=None, bold=False, italic=False,
                           align=None):
    """Insert a new paragraph right after the given one and return it."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        run = new_para.add_run(text)
        run.bold = bold
        run.italic = italic
        # Inherit Times New Roman 14pt
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
    if align == "center":
        new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return new_para


def insert_image_after(paragraph, image_path, width_inches=6.5):
    """Insert a centered image paragraph after the given paragraph."""
    p = insert_paragraph_after(paragraph, "", align="center")
    run = p.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(width_inches))
    except Exception as e:
        print(f"  ! could not embed {image_path}: {e}")
    return p


def find_paragraph_index(doc, predicate):
    """Find first paragraph index whose text starts with prefix (or matches predicate)."""
    for i, p in enumerate(doc.paragraphs):
        if predicate(p):
            return i
    return -1


def find_heading(doc, level, text_prefix):
    """Find paragraph by heading level and text prefix. Returns index or -1."""
    target = f"Heading {level}"
    text_norm = text_prefix.strip().lower()
    for i, p in enumerate(doc.paragraphs):
        sn = (p.style.name if p.style else "").lower()
        if target.lower() in sn and p.text.strip().lower().startswith(text_norm):
            return i
    return -1


def find_next_heading(doc, start_idx, levels=("Heading 1", "Heading 2")):
    """Find index of the next heading at one of the given levels, after start_idx."""
    for i in range(start_idx + 1, len(doc.paragraphs)):
        sn = doc.paragraphs[i].style.name if doc.paragraphs[i].style else ""
        if sn in levels:
            return i
    return len(doc.paragraphs)


def delete_paragraphs_between(doc, start_exclusive, end_exclusive):
    """Delete paragraphs (start, end) exclusive on both ends."""
    paragraphs = doc.paragraphs
    to_delete = []
    for i in range(start_exclusive + 1, end_exclusive):
        to_delete.append(paragraphs[i]._p)
    for p in to_delete:
        p.getparent().remove(p)


def add_paragraphs_after(start_paragraph, blocks):
    """Insert a sequence of blocks after start_paragraph.

    blocks is a list of dicts:
      {"text": str, "bold": False, "italic": False, "align": "justify"}
      {"image": Path, "caption": str, "width": 6.5}
      {"heading": str, "level": 2}
    Returns the last inserted paragraph (for chaining).
    """
    cur = start_paragraph
    for b in blocks:
        if "image" in b:
            cur = insert_image_after(cur, b["image"], width_inches=b.get("width", 6.5))
            if b.get("caption"):
                cur = insert_paragraph_after(cur, b["caption"],
                                             italic=True, align="center")
        elif "heading" in b:
            cur = insert_paragraph_after(cur, b["heading"], align="justify")
            try:
                cur.style = f"Heading {b.get('level', 2)}"
            except Exception:
                pass
        else:
            cur = insert_paragraph_after(
                cur,
                b.get("text", ""),
                bold=b.get("bold", False),
                italic=b.get("italic", False),
                align=b.get("align", "justify"),
            )
    return cur


def marker_present(doc, marker_text):
    """Check if any paragraph contains a marker tag."""
    for p in doc.paragraphs:
        if marker_text in p.text:
            return True
    return False


# ---------------------------------------------------------------------------
# Operation definitions (text content)
# ---------------------------------------------------------------------------

TISSUE_HEATMAP_PARA = (
    "Распределение классов ядер существенно различается между тканями: "
    "в опухолевых тканях (Breast, Prostate, Liver, Ovarian) доминируют "
    "неопластические ядра, в иммунологически активных тканях (Stomach, Skin) — "
    "воспалительные, в Pancreatic и Lung — соединительнотканные. Класс Dead "
    "практически отсутствует во всех тканях кроме Lung (24%) и Thyroid (35%), "
    "что объясняется специфическими патологическими процессами. Эта тканевая "
    "специфичность создаёт дополнительный вызов для обучения единой модели: "
    "распределение классов в обучающем батче существенно зависит от того, какие "
    "ткани в нём представлены."
)
TISSUE_HEATMAP_CAPTION = (
    "Рисунок 1.4 — Распределение классов ядер по тканям PanNuke (% от ядер в "
    "данной ткани). Подтверждает существенную тканево-классовую "
    "неоднородность датасета и объясняет, почему наивный random sampling без "
    "учёта тканей даёт нестабильный обучающий сигнал на редких классах."
)

NULITE_HOVER_MAMBA = [
    (
        "NuLite. Tommasino et al. предложили лёгкую модель NuLite, "
        "ориентированную на одновременную сегментацию и классификацию ядер при "
        "существенно меньших вычислительных затратах. NuLite использует кодер "
        "FastViT — гибридную свёрточно-трансформерную архитектуру со структурной "
        "репараметризацией, — и один общий декодер HoVer-Net-стиля с тремя "
        "выходными головами. На бенчмарке PanNuke NuLite-T (12M параметров) "
        "достигает mPQ ≈ 0,47–0,49, NuLite-H (34M) — около 0,496. Сопоставимое "
        "качество при многократно меньшем размере получено за счёт архитектурной "
        "эффективности FastViT и тщательного training-рецепта (Focal Tversky Loss, "
        "вспомогательная классификация ткани, class-balanced sampling). NuLite не "
        "использует дистилляцию знаний и служит в настоящей работе главным "
        "внешним baseline'ом для сравнения лёгких моделей."
    ),
    (
        "HoVer-UNet. Tommasino et al. в более ранней работе применили "
        "response-based дистилляцию из CellViT-SAM-H в ученика с кодером "
        "ConvNeXt и декодером HoVer-Net, получив mPQ ≈ 0,49 на PanNuke. Это "
        "единственная прежняя работа, явно применяющая knowledge distillation к "
        "задаче сегментации клеточных ядер, и она используется в настоящей работе "
        "как методологический ориентир."
    ),
    (
        "Архитектуры на основе моделей пространства состояний (Mamba). "
        "В 2024–2025 годах появился новый класс архитектур для компьютерного "
        "зрения, основанный на selective state-space models (SSM): Mamba, Vision "
        "Mamba, VMamba. Mamba-блоки обеспечивают глобальное receptive field "
        "подобно трансформерам, но с линейной (а не квадратичной) сложностью по "
        "числу токенов. Для медицинской сегментации были предложены адаптации: "
        "VM-UNet, U-Mamba, UltraLight VM-UNet — последняя использует Parallel "
        "Vision Mamba (PVM) блок, разбивающий входные каналы на группы и "
        "применяющий независимый Mamba к каждой группе, что радикально сокращает "
        "число параметров. В области патологии CP-Mamba (AAAI 2025) — первая "
        "опубликованная Mamba-модель для сегментации ядер на PanNuke с "
        "использованием category-prompt supervision, достигает mPQ ≈ 0,615. "
        "Однако исследования по применению Mamba именно в качестве декодера в "
        "сочетании с дистилляцией знаний от свёрточно-трансформерного учителя в "
        "литературе отсутствуют — эта ниша мотивирует один из экспериментов "
        "настоящей работы."
    ),
]

PIPELINE_INTRO_PARA = (
    "Выделяют три основных типа дистилляции: response-based (на уровне выходов "
    "модели), feature-based (на уровне промежуточных представлений) и "
    "relation-based (на уровне отношений между представлениями). В настоящей "
    "работе подробно изучается response-based дистилляция в трёх различных "
    "формулировках функции потерь (разделы 2.4.1–2.4.3); feature-based "
    "дистилляция реализована в пайплайне как опциональная конфигурация, но не "
    "входит в основной сравнительный эксперимент. Общая архитектура потока "
    "обучения с предрасчитанными мягкими целями приведена на рис. 2.1."
)
PIPELINE_CAPTION = (
    "Рисунок 2.1 — Сквозной пайплайн обучения с дистилляцией знаний. Учитель "
    "прогоняется по датасету один раз оффлайн, и его логиты сохраняются на диск. "
    "При обучении ученика мягкие цели загружаются с диска, синхронно проходят "
    "через те же пространственные аугментации, что и входные изображения, и "
    "поступают в дистилляционный лосс одновременно с GT-целями."
)

STUDENT_ARCH_NEW = [
    (
        "Модель-ученик представляет собой кодер–декодер архитектуру с тремя "
        "выходными головами, аналогичными CellViT-SAM-H, но с существенно меньшим "
        "числом параметров."
    ),
    (
        "В качестве кодера используется FastViT-S12 — гибридная архитектура от "
        "Apple, сочетающая свёрточные блоки и self-attention с использованием "
        "structural reparameterization. На этапе обучения блок состоит из "
        "нескольких параллельных свёрточных ветвей, которые на этапе инференса "
        "математически сливаются в одну эквивалентную свёртку — это даёт "
        "«бесплатное» сокращение числа параметров и FLOPs во время инференса при "
        "сохранении ёмкости при обучении. FastViT-S12 содержит около 8,3 "
        "миллионов параметров и извлекает признаки на четырёх масштабных уровнях "
        "с числом каналов 64, 128, 256 и 512. Кодер инициализирован весами, "
        "предобученными на ImageNet-1k."
    ),
    (
        "Выбор FastViT-S12 мотивирован двумя соображениями. Во-первых, эта "
        "архитектура использована в NuLite и доказала способность достигать "
        "качества, сопоставимого с CellViT-SAM-H, при 50-кратной разнице в "
        "размере. Во-вторых, гибридная природа FastViT обеспечивает разумный "
        "компромисс между локальной природой свёрток (важной для морфологии "
        "ядер) и глобальным контекстом self-attention (важным для классификации "
        "типа ядра по тканевому контексту)."
    ),
    (
        "В качестве декодера в основной конфигурации используется FPN (Feature "
        "Pyramid Network) в стиле HoVer-Net: последовательные блоки выполняют "
        "билинейную интерполяцию для увеличения разрешения, конкатенацию с "
        "признаками соответствующего уровня кодера через skip connections и два "
        "свёрточных блока (Conv2d + BatchNorm + ReLU). Декодер содержит четыре "
        "уровня с числом каналов 256, 128, 64 и 32, восстанавливая разрешение до "
        "исходного 256×256. В рамках сравнительного исследования также реализован "
        "альтернативный декодер на основе блоков Mamba (раздел 2.6)."
    ),
    (
        "На выходе декодера располагаются три независимые головы: бинарная "
        "голова (2 канала — фон и ядро), HV-голова (2 канала — горизонтальная и "
        "вертикальная карты расстояний, опционально с tanh-активацией для "
        "ограничения диапазона [−1, 1]) и голова классификации типов (6 каналов — "
        "фон и 5 классов ядер). Каждая голова состоит из одного блока ConvBNReLU "
        "и свёртки 1×1. Дополнительно реализована вспомогательная голова "
        "классификации ткани (19 классов PanNuke), подключаемая к выходу самой "
        "глубокой стадии кодера через global average pooling и линейный слой; "
        "эта голова используется для multi-task регуляризации обучения."
    ),
    (
        "Суммарно базовая модель-ученик содержит около 11,5 миллионов "
        "параметров — в 55 раз меньше, чем CellViT-SAM-H, и сопоставимо по "
        "размеру с NuLite-T (12M)."
    ),
]

LOSS_INTRO = (
    "Итоговая функция потерь объединяет компоненты supervised loss (на основе "
    "ground truth разметки) и дистилляционные компоненты (на основе мягких целей "
    "учителя): L = L_GT + α · L_KD, где α — гиперпараметр баланса. Supervised "
    "loss остаётся одинаковым во всех экспериментах: L_GT = λ₁·L_FTL(binary) + "
    "λ₂·L_Dice(binary) + λ₃·L_MSE(hv) + λ₄·L_MSGE(hv) + λ₅·L_Focal(type) + "
    "λ₆·L_Dice(type) + λ₇·L_CE(tissue)."
)
LOSS_DETAILS = (
    "Для бинарной головы используется Focal Tversky Loss (Abraham & Khan, 2019), "
    "более чувствительный к ложно-отрицательным предсказаниям на мелких объектах "
    "(что важно для PanNuke, где медианный размер ядра — несколько десятков "
    "пикселей). Для HV-головы — MSE в маске ядер и mean squared gradient error "
    "(MSGE) на Sobel-градиентах: MSGE штрафует размытие границ между соседними "
    "ядрами, где градиенты HV-карт наиболее информативны для алгоритма "
    "водораздела. Для головы типов — Focal Loss (Lin et al., 2017) с обратно-"
    "частотными весами классов, частично компенсирующими дисбаланс. "
    "Вспомогательная голова ткани обучается обычной кросс-энтропией с малым "
    "весом (λ₇ = 0,1)."
)
LOSS_KD_INTRO = (
    "Дистилляционный компонент L_KD применяется только к классификационным "
    "головам (binary, type), но не к регрессионной HV-голове. В работе "
    "реализованы и сравниваются три варианта формулировки L_KD, описанные в "
    "разделах 2.4.1–2.4.3. Концептуальное различие между ними проиллюстрировано "
    "на рис. 2.2."
)
KD_VARIANTS_CAPTION = (
    "Рисунок 2.2 — Три варианта response-based дистилляции, сравниваемые в "
    "работе: классическая KL-дивергенция (Hinton et al., 2015), Decoupled KD "
    "(Zhao et al., 2022) и Frequency-Decoupled KD (Lu et al., 2025). Различия — "
    "в том, какой компонент сигнала учителя усиливается: уровень классов, "
    "разделение target/non-target или частотные компоненты пространственной "
    "карты."
)

KL_SECTION = [
    {"heading": "2.4.1 KL-дивергенция (классическая дистилляция Hinton)", "level": 3},
    {"text":
        "Для классификационной головы с выходным вектором логитов z дистилляция "
        "определяется через softmax с температурой T: σ(z_i, T) = exp(z_i/T) / "
        "Σ_j exp(z_j/T). При T = 1 получается стандартный softmax; при T > 1 "
        "распределение становится более «мягким», усиливая информацию о "
        "соотношении между классами. Базовый дистилляционный loss определяется "
        "как KL-дивергенция между смягчёнными выходами student и teacher, "
        "усреднённая по всем пространственным позициям: L_KL = T² · KL(σ(z_t, T) "
        "|| σ(z_s, T)). Множитель T² компенсирует уменьшение градиентов при "
        "высоких температурах. В настоящей работе используется T = 10 — на "
        "порядок выше стандартных T = 4 из оригинальной работы Hinton, "
        "мотивировано экстремальным диапазоном выходных логитов CellViT-SAM-H "
        "(от −44 до +45), при котором softmax даже при T = 4 остаётся почти "
        "one-hot."},
]

DKD_SECTION = [
    {"heading": "2.4.2 Decoupled KD (Zhao et al., CVPR 2022)", "level": 3},
    {"text":
        "Анализ, выполненный авторами Decoupled KD, показывает, что стандартная "
        "KL-дивергенция допускает декомпозицию на две комплементарные "
        "составляющие — Target-Class KD (TCKD) и Non-target-Class KD (NCKD): "
        "L_KL = L_TCKD + (1 − p_t(c)) · L_NCKD, где c — целевой класс, p_t(c) — "
        "уверенность учителя в нём. TCKD представляет собой бинарную "
        "KL-дивергенцию между (p(c), 1 − p(c)) распределениями студента и "
        "учителя — измеряет согласие в уверенности на правильном классе. NCKD — "
        "KL-дивергенция по не-target классам после ренормализации — измеряет "
        "согласие в относительном ранжировании неправильных классов и несёт "
        "основную долю «тёмного знания»."},
    {"text":
        "Ключевое наблюдение: NCKD автоматически подавляется множителем "
        "(1 − p_t(c)), который стремится к нулю когда учитель уверен в целевом "
        "классе. Для редких классов (Dead в PanNuke присутствует в <2% патчей), "
        "где учитель часто очень уверен в фоновом классе, NCKD-сигнал "
        "теряется, и студент не получает информации о различении редких классов. "
        "DKD устраняет это подавление через независимые веса: L_DKD = α · L_TCKD "
        "+ β · L_NCKD, с рекомендуемыми авторами значениями α = 1, β = 8. В "
        "настоящей работе DKD применяется для исправления одного из неудачных "
        "режимов vanilla-KL на классе Dead, наблюдаемого в предварительных "
        "экспериментах. Визуальная иллюстрация декомпозиции приведена на "
        "рис. 2.3."},
]
DKD_CAPTION = (
    "Рисунок 2.3 — Декомпозиция полной KL-дивергенции на TCKD и NCKD по Zhao et "
    "al. (CVPR 2022). В правом блоке — DKD с β = 8, амплифицирующим NCKD-сигнал "
    "для борьбы с подавлением «тёмного знания» на редких классах."
)
DKD_POST = (
    "Для адаптации DKD к задаче плотного предсказания (сегментации) в качестве "
    "target-класса на каждом пикселе используется argmax-предсказание учителя, "
    "что сохраняет сигнатуру дистилляционного лосса без необходимости в ground "
    "truth разметке."
)

UFD_SECTION = [
    {"heading": "2.4.3 Frequency-Decoupled KD (Lu et al., BMVC 2025) — negative result", "level": 3},
    {"text":
        "Frequency-Decoupled Distillation предлагает разложение дистилляционного "
        "сигнала не по классам, а по пространственным частотам. Идея: на "
        "softmax-карте сегментации низкочастотные (LF) компоненты несут "
        "глобальную классовую структуру изображения, а высокочастотные (HF) — "
        "информацию о границах объектов и мелких структурах (включая редкие "
        "классы). Авторы предлагают применять двумерное дискретное косинусное "
        "преобразование (DCT-II) к softmax-картам учителя и студента, разделять "
        "коэффициенты на LF-блок (топ-левый K×K) и HF-остаток, и применять "
        "взвешенный MSE на каждой полосе: L_UFD = w_LF · MSE(s_LF, t_LF) + "
        "w_HF · MSE(s_HF, t_HF). При w_HF > w_LF метод предположительно "
        "усиливает сигнал на мелких/редких объектах."},
    {"text":
        "Эта схема была реализована в данной работе и применена к задаче "
        "сегментации ядер. Эмпирический анализ показал, что магнитуда "
        "LF-компонент softmax-карт превосходит HF-компонент примерно в 600 раз "
        "(LF ≈ 0,4 vs HF ≈ 7·10⁻⁴ в типичном батче), что делает «частотное "
        "разделение» при любых разумных значениях w_HF численно близким к "
        "чистому LF-MSE. То есть на задаче плотной сегментации с сильно "
        "сегментированными softmax-распределениями частотное декомплексирование "
        "вырождается. Метод сохранён в текущей работе как сравнительный baseline "
        "и явно отмечается как negative result для honest reporting."},
]

MAMBA_SECTION = [
    {"heading": "2.6 Альтернативный декодер на основе Mamba", "level": 2},
    {"text":
        "В дополнение к свёрточному декодеру в стиле HoVer-Net (раздел 2.3) в "
        "работе реализован альтернативный декодер на основе selective "
        "state-space модели Mamba. Mamba обеспечивает глобальное receptive field "
        "подобно self-attention, но с линейной по числу токенов сложностью, что "
        "делает её привлекательной для плотных задач компьютерного зрения с "
        "большим количеством пространственных позиций."},
    {"text":
        "Реализация следует подходу UltraLight VM-UNet — Parallel Vision Mamba "
        "(PVM). Входные каналы каждого декодерного уровня разделяются на "
        "G = 4 группы, каждая обрабатывается независимым Mamba-блоком, после "
        "чего конкатенируется и пропускается через GroupNorm с residual "
        "соединением. Такая факторизация радикально снижает число параметров по "
        "сравнению с применением одного Mamba-блока к полной размерности "
        "каналов: для канала C группа из четырёх Mamba даёт 4 · (C/4)² = C²/4 "
        "параметров в проекциях вместо C². Для скана пространственных позиций "
        "используется 4-way cross-scan (как в VMamba) — последовательность "
        "токенов проходится в четырёх направлениях (слева→направо, "
        "справа→налево, сверху→вниз, снизу→вверх) с усреднением результатов. "
        "Структурная аналогия со свёрточным декодером показана на рис. 2.4."},
]
MAMBA_CAPTION = (
    "Рисунок 2.4 — Сравнение свёрточного декодера в стиле HoVer-Net (слева) и "
    "Mamba-декодера в стиле UltraLight VM-UNet (справа). Кодерные стадии "
    "одинаковы (FastViT-S12); различия — в способе обработки признаков на каждом "
    "декодерном уровне."
)
MAMBA_POST = (
    "Применение Mamba именно в декодере при сохранении свёрточно-"
    "трансформерного кодера и в комбинации с дистилляцией знаний в литературе "
    "не описано. Это формирует одну из защищаемых научных нив настоящей работы."
)

ABLATION_GRID_TEXT = [
    "Сравнительное исследование организовано как двухфакторная ablation-сетка "
    "по типу декодера и формулировке дистилляционного сигнала. Каждая ячейка "
    "повторяется на трёх фолдах PanNuke (стандартный 3-fold cross-validation "
    "протокол) для оценки дисперсии. Условия: HoVer-Net декодер и Mamba "
    "декодер; KD-сигналы: без KD (baseline), классическая KL-дистилляция, "
    "Decoupled KD. Итого 2 × 3 × 3 = 18 обучающих запусков.",
    "Эта схема позволяет одновременно оценить: (1) вклад дистилляции при "
    "фиксированной архитектуре — горизонтальное сравнение в каждой строке "
    "ablation-таблицы; (2) сравнение декодеров при одинаковом обучающем "
    "сигнале — вертикальное сравнение в каждом столбце; (3) преимущество DKD "
    "над KL на редких классах — KL vs DKD; (4) интеракцию декодер × KD-сигнал — "
    "даёт ли Mamba преимущество именно с дистилляцией, или независимо.",
    "Дополнительные эксперименты вне основной сетки: Hybrid A — оценка "
    "предобученного чекпоинта NuLite-T на нашем оценочном пайплайне для "
    "прямого external сравнения. Hybrid B — обучение NuLite-T архитектуры на "
    "нашем training-рецепте для изоляции эффекта recipe от архитектуры. "
    "Test-time augmentation — 8-way TTA (4 поворота × 2 отражения с корректной "
    "sign-коррекцией HV-компонент) применяется ко всем основным конфигурациям "
    "после обучения. MoNuSeg cross-dataset eval — zero-shot оценка лучшего "
    "student на внешнем датасете MoNuSeg для оценки domain-shift устойчивости.",
    "Все эксперименты оцениваются стандартными метриками: multi-class Panoptic "
    "Quality (mPQ), binary Panoptic Quality (bPQ) и F1-score детекции при "
    "IoU ≥ 0,5. Для оценки статистической значимости различий между "
    "конфигурациями используется парный критерий Уилкоксона на per-image PQ "
    "массивах и bootstrap-доверительные интервалы.",
]

NOVELTY_NEW = [
    "В работе получены следующие результаты, обладающие научной новизной:",
    "• Сравнительное исследование KD-сигналов для сегментации клеточных "
    "ядер. Впервые на единой архитектуре ученика и едином training-рецепте "
    "систематически сравниваются три формулировки response-based дистилляции "
    "(классическая KL, Decoupled KD, Frequency-Decoupled KD) с организацией "
    "экспериментов по схеме ablation-сетки 2 декодера × 3 KD-сигнала × 3 фолда.",
    "• Применение Mamba (selective state-space модели) в качестве декодера в "
    "задаче сегментации клеточных ядер в комбинации с дистилляцией знаний. "
    "Ранее опубликованные применения Mamba к этой задаче (CP-Mamba, AAAI 2025) "
    "использовали category-prompt supervision, но не дистилляцию; работы по "
    "дистилляции в Mamba существуют только для других task family (LLM, "
    "super-resolution, semantic segmentation outdoor scenes).",
    "• Обнаружение и устранение трёх методологических дефектов, имеющих "
    "существенное влияние на корректность результатов дистилляции для dense "
    "prediction: (1) пространственная рассогласованность предрасчитанных "
    "мягких целей с аугментированным входом ученика (теряется ~87% обучающего "
    "сигнала); (2) некорректный протокол подсчёта mPQ в условиях дисбаланса "
    "классов (систематическое занижение в 2,4 раза); (3) чувствительность KD к "
    "выбору α в зависимости от среды исполнения (PyTorch версии и batch size). "
    "Каждая находка задокументирована и сопровождается фиксом в исходном коде.",
    "• Honest negative result для Frequency-Decoupled KD на задаче плотной "
    "сегментации: показано, что LF-компонент softmax-карт превосходит "
    "HF-компонент по магнитуде примерно в 600 раз, что делает «частотное "
    "декомплексирование» в данной постановке вырожденным независимо от выбора "
    "весов.",
]

PRACTICAL_NEW = (
    "Лучшая из обученных моделей (FastViT-S12 со свёрточным декодером + "
    "response-based KD + 8-way TTA) достигает mPQ ≈ 0,47 на 3-фолдовой "
    "кросс-валидации PanNuke при 11,5 миллионах параметров, что соответствует "
    "80 % качества 630-миллионной модели-учителя CellViT-SAM-H в едином "
    "пайплайне оценки. Сокращение размера в 55 раз и пиковой видеопамяти "
    "инференса с ≥ 24 ГБ до ~2 ГБ делает развёртывание возможным на "
    "потребительских GPU (16 ГБ и менее), что расширяет возможности применения "
    "автоматической сегментации клеточных ядер в клинических учреждениях, не "
    "располагающих дорогостоящим графическим оборудованием."
)

# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if not DOCX_PATH.exists():
        print(f"ERROR: {DOCX_PATH} not found")
        sys.exit(1)

    # Backup
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)
        print(f"Backup created at {BACKUP_PATH}")

    doc = Document(str(DOCX_PATH))

    # Idempotency guard
    if marker_present(doc, MARKER):
        print(f"Marker {MARKER} already in document — script already applied. "
              "Restore from .bak.docx to re-run.")
        return

    print("\n--- Op 1: §1.2 add tissue heatmap (text + figure 1.4) ---")
    idx = find_paragraph_index(
        doc, lambda p: "распределение патчей по 19 типам тканей" in p.text.lower()
    )
    if idx == -1:
        # fall back: end of §1.2 paragraph mentioning "Рисунок 1.3"
        idx = find_paragraph_index(
            doc, lambda p: "рисунок 1.3" in p.text.lower()
        )
    if idx > 0:
        anchor = doc.paragraphs[idx]
        last = add_paragraphs_after(anchor, [
            {"text": TISSUE_HEATMAP_PARA},
            {"image": FIG_DIR / "10_per_tissue_class_heatmap.png", "width": 6.3,
             "caption": TISSUE_HEATMAP_CAPTION},
        ])
        print(f"  ✓ inserted after para {idx}")
    else:
        print("  ! anchor for §1.2 heatmap not found, skipping")

    print("\n--- Op 2: §1.4 append NuLite/HoVer-UNet/Mamba paragraphs ---")
    # Find §1.4 heading then last paragraph before next chapter
    h_idx = find_heading(doc, 2, "1.4 Обзор архитектур")
    if h_idx > 0:
        next_idx = find_next_heading(doc, h_idx, ("Heading 1",))
        anchor = doc.paragraphs[next_idx - 1]
        add_paragraphs_after(anchor, [{"text": t} for t in NULITE_HOVER_MAMBA])
        print(f"  ✓ appended 3 paragraphs after para {next_idx - 1}")

    print("\n--- Op 3: §2.1 add KD-variants intro + pipeline figure 2.1 ---")
    h_idx = find_heading(doc, 2, "2.1 Дистилляция знаний")
    if h_idx > 0:
        # Find end of §2.1 (next Heading 2)
        next_idx = find_next_heading(doc, h_idx, ("Heading 2", "Heading 1"))
        anchor = doc.paragraphs[next_idx - 1]
        add_paragraphs_after(anchor, [
            {"text": PIPELINE_INTRO_PARA},
            {"image": FIG_DIR / "08_pipeline_diagram.png", "width": 6.5,
             "caption": PIPELINE_CAPTION},
        ])
        print(f"  ✓ inserted at end of §2.1 (para {next_idx - 1})")

    print("\n--- Op 4: §2.3 rewrite student architecture (FastViT) ---")
    h_idx = find_heading(doc, 2, "2.3 Архитектура модели-ученика")
    if h_idx > 0:
        next_idx = find_next_heading(doc, h_idx, ("Heading 2", "Heading 1"))
        # Delete existing §2.3 content
        delete_paragraphs_between(doc, h_idx, next_idx)
        # Re-find positions after deletion
        h_idx = find_heading(doc, 2, "2.3 Архитектура модели-ученика")
        anchor = doc.paragraphs[h_idx]
        add_paragraphs_after(anchor, [{"text": t} for t in STUDENT_ARCH_NEW])
        print(f"  ✓ rewrote §2.3 ({len(STUDENT_ARCH_NEW)} paragraphs)")

    print("\n--- Op 5: §2.4 replace with KL/DKD/UFD subsections ---")
    h_idx = find_heading(doc, 2, "2.4 Функция потерь")
    if h_idx > 0:
        next_idx = find_next_heading(doc, h_idx, ("Heading 2", "Heading 1"))
        delete_paragraphs_between(doc, h_idx, next_idx)
        h_idx = find_heading(doc, 2, "2.4 Функция потерь")
        # Also rename heading? Keep "2.4 Функции потерь" — change one letter to plural
        doc.paragraphs[h_idx].runs[0].text = "2.4 Функции потерь"
        anchor = doc.paragraphs[h_idx]
        blocks = [
            {"text": LOSS_INTRO},
            {"text": LOSS_DETAILS},
            {"text": LOSS_KD_INTRO},
            {"image": FIG_DIR / "09_kd_variants.png", "width": 6.5,
             "caption": KD_VARIANTS_CAPTION},
        ] + KL_SECTION + DKD_SECTION + [
            {"image": FIG_DIR / "12_dkd_decomposition.png", "width": 6.3,
             "caption": DKD_CAPTION},
            {"text": DKD_POST},
        ] + UFD_SECTION
        add_paragraphs_after(anchor, blocks)
        print(f"  ✓ §2.4 fully replaced with {len(blocks)} blocks")

    print("\n--- Op 6: insert NEW §2.6 Mamba decoder before old §2.6 ---")
    # Old §2.6 is "Дизайн экспериментов"; new §2.6 is "Альтернативный декодер на основе Mamba"
    h_idx = find_heading(doc, 2, "2.6 Дизайн экспериментов")
    if h_idx > 0:
        # Rename old §2.6 to §2.7 first
        doc.paragraphs[h_idx].runs[0].text = "2.7 Дизайн экспериментов"
        # Now insert new §2.6 BEFORE this (i.e., after preceding paragraph)
        # Find last paragraph of §2.5
        prev_anchor = doc.paragraphs[h_idx - 1]
        blocks = MAMBA_SECTION + [
            {"image": FIG_DIR / "11_decoder_comparison.png", "width": 6.5,
             "caption": MAMBA_CAPTION},
            {"text": MAMBA_POST},
        ]
        add_paragraphs_after(prev_anchor, blocks)
        print(f"  ✓ inserted new §2.6 (Mamba) and renamed old →§2.7")

    print("\n--- Op 7: §2.7 rewrite design-of-experiments (ablation grid) ---")
    h_idx = find_heading(doc, 2, "2.7 Дизайн экспериментов")
    if h_idx > 0:
        next_idx = find_next_heading(doc, h_idx, ("Heading 2", "Heading 1"))
        delete_paragraphs_between(doc, h_idx, next_idx)
        h_idx = find_heading(doc, 2, "2.7 Дизайн экспериментов")
        anchor = doc.paragraphs[h_idx]
        add_paragraphs_after(anchor, [{"text": t} for t in ABLATION_GRID_TEXT])
        print(f"  ✓ §2.7 rewritten with {len(ABLATION_GRID_TEXT)} paragraphs")

    print("\n--- Op 8: Введение → Научная новизна rewrite ---")
    h_idx = find_paragraph_index(
        doc, lambda p: p.text.strip() == "Научная новизна"
    )
    if h_idx > 0:
        next_idx = find_next_heading(doc, h_idx, ("Heading 1", "Heading 2"))
        # Also handle if next title under same heading 2 like "Практическая значимость"
        for j in range(h_idx + 1, len(doc.paragraphs)):
            if "Практическая значимость" in doc.paragraphs[j].text or \
               "Структура работы" in doc.paragraphs[j].text:
                next_idx = j
                break
        delete_paragraphs_between(doc, h_idx, next_idx)
        h_idx = find_paragraph_index(doc, lambda p: p.text.strip() == "Научная новизна")
        anchor = doc.paragraphs[h_idx]
        add_paragraphs_after(anchor, [{"text": t} for t in NOVELTY_NEW])
        print(f"  ✓ rewrote Научная новизна ({len(NOVELTY_NEW)} paragraphs)")

    print("\n--- Op 9: Введение → Практическая значимость rewrite ---")
    h_idx = find_paragraph_index(
        doc, lambda p: p.text.strip() == "Практическая значимость"
    )
    if h_idx > 0:
        next_idx = find_next_heading(doc, h_idx, ("Heading 1", "Heading 2"))
        for j in range(h_idx + 1, len(doc.paragraphs)):
            if "Структура работы" in doc.paragraphs[j].text:
                next_idx = j
                break
        delete_paragraphs_between(doc, h_idx, next_idx)
        h_idx = find_paragraph_index(doc, lambda p: p.text.strip() == "Практическая значимость")
        anchor = doc.paragraphs[h_idx]
        add_paragraphs_after(anchor, [{"text": PRACTICAL_NEW}])
        print("  ✓ rewrote Практическая значимость")

    # Add invisible marker at the very end so re-runs detect and skip.
    end_p = doc.add_paragraph(MARKER)
    end_p.runs[0].font.size = Pt(1)
    end_p.runs[0].font.color.rgb = None
    # render as white-ish to make invisible
    from docx.shared import RGBColor
    end_p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.save(str(DOCX_PATH))
    print(f"\nSaved updated docx: {DOCX_PATH}")
    print(f"Backup at: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
